from dotenv import load_dotenv
load_dotenv()

import os
import base64
import json
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from docx import Document
from fastapi import FastAPI, HTTPException, Security, Depends
from fastapi.security.api_key import APIKeyHeader
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from io import BytesIO
import logging
import httpx

# Windows Tesseract path
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="AI Document Analysis API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

API_KEY = os.getenv("API_KEY", "sk_track2_987654321")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

api_key_header = APIKeyHeader(name="x-api-key", auto_error=False)


class DocumentRequest(BaseModel):
    fileName: str
    fileType: str
    fileBase64: str


def verify_api_key(api_key: str = Security(api_key_header)):
    if not api_key or api_key != API_KEY:
        raise HTTPException(status_code=401, detail="Invalid or missing API key")
    return api_key


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return "\n".join(text_parts).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n".join(paragraphs).strip()


def extract_text_from_image(file_bytes: bytes) -> str:
    image = Image.open(BytesIO(file_bytes))
    text = pytesseract.image_to_string(image, config="--psm 6")
    return text.strip()


def analyze_with_gemini(text: str, file_name: str) -> dict:
    prompt = f"""You are a document analysis AI. Analyze the following document text and return a JSON response.

Document: {file_name}

Text:
{text[:8000]}

Return ONLY a valid JSON object (no markdown, no extra text) with this exact structure:
{{
  "summary": "A concise 1-3 sentence summary of the document content",
  "entities": {{
    "names": ["list of person names found"],
    "dates": ["list of dates found"],
    "organizations": ["list of organization/company names found"],
    "amounts": ["list of monetary amounts found"],
    "locations": ["list of locations/places found"]
  }},
  "sentiment": "Positive or Neutral or Negative"
}}

Rules:
- summary: Capture the core purpose and content of the document
- entities: Extract ALL named entities present. Use empty arrays if none found
- sentiment: Positive means good news/praise/success, Negative means complaints/issues/failures, Neutral means factual/informational
"""

    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={GEMINI_API_KEY}"

    payload = {
        "contents": [
            {
                "parts": [{"text": prompt}]
            }
        ]
    }

    response = httpx.post(url, json=payload, timeout=30)
    response.raise_for_status()

    data = response.json()
    raw = data["candidates"][0]["content"]["parts"][0]["text"].strip()

    # Strip markdown fences if present
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    raw = raw.strip()

    return json.loads(raw)


@app.get("/")
def root():
    return {"status": "ok", "message": "AI Document Analysis API is running"}


@app.get("/health")
def health():
    return {"status": "healthy"}


@app.post("/api/document-analyze")
async def analyze_document(
    request: DocumentRequest,
    api_key: str = Depends(verify_api_key)
):
    file_type = request.fileType.lower().strip()
    if file_type not in ["pdf", "docx", "image"]:
        raise HTTPException(status_code=400, detail="Unsupported fileType. Use pdf, docx, or image.")

    try:
        file_bytes = base64.b64decode(request.fileBase64)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid base64 encoded file.")

    try:
        if file_type == "pdf":
            extracted_text = extract_text_from_pdf(file_bytes)
        elif file_type == "docx":
            extracted_text = extract_text_from_docx(file_bytes)
        elif file_type == "image":
            extracted_text = extract_text_from_image(file_bytes)
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Failed to extract text: {str(e)}")

    if not extracted_text:
        raise HTTPException(status_code=422, detail="No text could be extracted.")

    try:
        analysis = analyze_with_gemini(extracted_text, request.fileName)
    except Exception as e:
        logger.error(f"Gemini error: {e}")
        raise HTTPException(status_code=500, detail=f"AI analysis failed: {str(e)}")

    return {
        "status": "success",
        "fileName": request.fileName,
        "summary": analysis.get("summary", ""),
        "entities": analysis.get("entities", {
            "names": [], "dates": [],
            "organizations": [], "amounts": [], "locations": []
        }),
        "sentiment": analysis.get("sentiment", "Neutral")
    }