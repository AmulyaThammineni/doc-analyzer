import asyncio
import base64
import json
import logging
import os
from io import BytesIO
from typing import Any, Dict, Optional

import fitz
import httpx
import pytesseract
from docx import Document
from PIL import Image

logger = logging.getLogger(__name__)


def configure_tesseract() -> None:
    """
    Configure pytesseract in a cross-platform way.

    - If TESSERACT_CMD is set, use it.
    - Otherwise, rely on Tesseract being available on PATH.
    """
    cmd = os.environ.get("TESSERACT_CMD")
    if cmd:
        pytesseract.pytesseract.tesseract_cmd = cmd


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            text_parts.append(page.get_text())
    return "\n".join(text_parts).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n".join(paragraphs).strip()


def extract_text_from_image(file_bytes: bytes) -> str:
    configure_tesseract()
    image = Image.open(BytesIO(file_bytes))
    text = pytesseract.image_to_string(image, config="--psm 6")
    return text.strip()


def detect_file_type(file_name: str) -> str:
    name = (file_name or "").lower().strip()
    if name.endswith(".pdf"):
        return "pdf"
    if name.endswith(".docx"):
        return "docx"
    if name.endswith((".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff")):
        return "image"
    return ""


def extract_text(file_bytes: bytes, file_type: str) -> str:
    t = (file_type or "").lower().strip()
    if t == "pdf":
        return extract_text_from_pdf(file_bytes)
    if t == "docx":
        return extract_text_from_docx(file_bytes)
    if t == "image":
        return extract_text_from_image(file_bytes)
    raise ValueError("Use pdf, docx, or image")


async def analyze_with_gemini(text: str, file_name: str, *, model: Optional[str] = None) -> Dict[str, Any]:
    prompt = f"""You are a document analysis AI. Analyze the following document text and return a JSON response.

Document: {file_name}

Text:
{text[:8000]}

Return ONLY a valid JSON object (no markdown, no extra text) with this exact structure:
{{
  "summary": "A concise 1-3 sentence summary",
  "entities": {{
    "names": [],
    "dates": [],
    "organizations": [],
    "amounts": [],
    "locations": []
  }},
  "sentiment": "Positive or Neutral or Negative"
}}
"""

    gemini_key = os.environ.get("GEMINI_API_KEY")
    use_model = model or os.environ.get("GEMINI_MODEL") or "gemini-2.5-flash"

    if not gemini_key:
        raise RuntimeError("GEMINI_API_KEY is missing")

    url = f"https://generativelanguage.googleapis.com/v1/models/{use_model}:generateContent"
    headers = {"Content-Type": "application/json", "x-goog-api-key": gemini_key}
    payload = {"contents": [{"parts": [{"text": prompt}]}]}

    try:
        async with httpx.AsyncClient() as client:
            response = await client.post(url, headers=headers, json=payload, timeout=60)
            response.raise_for_status()
    except httpx.HTTPStatusError as e:
        logger.error("Gemini API error: %s", e.response.text)
        raise RuntimeError(f"Gemini API failed: {e.response.text}") from e
    except Exception as e:
        logger.error("Request failed: %s", str(e))
        raise RuntimeError(f"Request failed: {str(e)}") from e

    data = response.json()
    try:
        raw = data["candidates"][0]["content"]["parts"][0]["text"].strip()
    except (KeyError, IndexError) as e:
        logger.error("Unexpected response: %s", data)
        raise RuntimeError("Invalid response format from AI") from e

    # Clean markdown formatting
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    raw = raw.strip()

    try:
        return json.loads(raw)
    except json.JSONDecodeError as e:
        logger.error("JSON parse failed. Raw: %s", raw)
        raise RuntimeError("Failed to parse AI response") from e


async def analyze_document_bytes(file_bytes: bytes, file_name: str, file_type: str) -> Dict[str, Any]:
    extracted_text = extract_text(file_bytes, file_type)
    if not extracted_text:
        raise ValueError("No text extracted")
    analysis = await analyze_with_gemini(extracted_text, file_name)
    return {
        "status": "success",
        "fileName": file_name,
        "summary": analysis.get("summary", ""),
        "entities": analysis.get(
            "entities",
            {"names": [], "dates": [], "organizations": [], "amounts": [], "locations": []},
        ),
        "sentiment": analysis.get("sentiment", "Neutral"),
        "extractedText": extracted_text,
    }


def analyze_document_bytes_sync(file_bytes: bytes, file_name: str, file_type: str) -> Dict[str, Any]:
    # Streamlit runs in a normal thread; this keeps the UI code simple.
    return asyncio.run(analyze_document_bytes(file_bytes, file_name, file_type))


def to_base64(file_bytes: bytes) -> str:
    return base64.b64encode(file_bytes).decode("utf-8")

